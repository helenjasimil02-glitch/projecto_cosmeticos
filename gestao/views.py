import json
import csv
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Produto, Venda, ItemVenda, ItemCompra, Compra, Fornecedor, Despesa, ReceitaExtra, Categoria
from django.db.models import Sum, F, Q
from django.utils import timezone
from datetime import timedelta
from django.db import transaction
from django.http import HttpResponse
from django.core.paginator import Paginator


@login_required
def dashboard(request):
    hoje = timezone.now().date()
    inicio_mes = hoje.replace(day=1)
    produtos = Produto.objects.all()

    # --- 1. KPIs DO MÊS ATUAL (DESEMPENHO) ---
    v_mes = float(Venda.objects.filter(data__gte=inicio_mes).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
    e_mes = float(ReceitaExtra.objects.filter(data__gte=inicio_mes).aggregate(Sum('valor'))['valor__sum'] or 0)
    c_mes = float(Compra.objects.filter(data__gte=inicio_mes).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
    d_mes = float(Despesa.objects.filter(data__gte=inicio_mes).aggregate(Sum('valor'))['valor__sum'] or 0)
    
    entradas_mes, saidas_mes = (v_mes + e_mes), (c_mes + d_mes)
    lucro_mes = entradas_mes - saidas_mes

    # --- 2. SALDO HISTÓRICO ACUMULADO ---
    total_vendas_geral = float(Venda.objects.aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
    total_extras_geral = float(ReceitaExtra.objects.aggregate(Sum('valor'))['valor__sum'] or 0)
    total_compras_geral = float(Compra.objects.aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
    total_despesas_geral = float(Despesa.objects.aggregate(Sum('valor'))['valor__sum'] or 0)
    
    saldo_caixa_real = (total_vendas_geral + total_extras_geral) - (total_compras_geral + total_despesas_geral)

    # --- 3. PATRIMÓNIO ATUAL ---
    valor_stock = sum(float(p.valor_total_stock()) for p in produtos)
    capital_giro_consolidado = saldo_caixa_real + valor_stock

    # --- 4. INTELIGÊNCIA E GRÁFICOS ---
    capital_estagnado = sum(float(p.valor_total_stock()) for p in produtos if p.status_giro() == "Estagnado ⚠️")
    
    count_a = len([p for p in produtos if p.classe_abc() == 'A' and p.stock_actual > 0])
    count_b = len([p for p in produtos if p.classe_abc() == 'B' and p.stock_actual > 0])
    count_c = len([p for p in produtos if p.classe_abc() == 'C' and p.stock_actual > 0])

    sugestao = [p for p in produtos if p.classe_abc() in ['A', 'B'] and p.stock_actual <= p.stock_minimo]

    vendas_diarias = []
    dias_semana = []
    for i in range(6, -1, -1):
        dia = hoje - timedelta(days=i)
        dias_semana.append(dia.strftime('%d/%m'))
        v = Venda.objects.filter(data__date=dia).aggregate(Sum('valor_total'))['valor_total__sum'] or 0
        vendas_diarias.append(float(v))

    context = {
        'entradas': entradas_mes, 'saidas': saidas_mes, 'lucro': lucro_mes,
        'valor_inventario': valor_stock, 
        'capital_giro': capital_giro_consolidado,
        'capital_estagnado': capital_estagnado,
        'sugestao_compra': sugestao,
        'abc_counts': [count_a, count_b, count_c],
        'vendas_diarias': vendas_diarias, 'dias_semana': dias_semana,
        'alertas_stock': Produto.objects.filter(stock_actual__lte=F('stock_minimo'))[:4],
        'alertas_validade': ItemCompra.objects.filter(validade__lte=hoje + timedelta(days=30), quantidade__gt=0).order_by('validade')[:4],
        'status_cor': 'text-success' if lucro_mes >= 0 else 'text-danger'
    }
    return render(request, 'dashboard.html', context)


@login_required
def registrar_venda(request):
    if request.method == "POST":
        carrinho_json = request.POST.get('carrinho_dados')

        # Validar se o carrinho não está vazio
        if not carrinho_json:
            messages.error(request, "O carrinho está vazio. Adiciona pelo menos um produto.")
            return redirect('registrar_venda')

        try:
            itens = json.loads(carrinho_json)
        except json.JSONDecodeError:
            messages.error(request, "Erro ao processar o carrinho. Tenta novamente.")
            return redirect('registrar_venda')

        if not itens:
            messages.error(request, "O carrinho está vazio. Adiciona pelo menos um produto.")
            return redirect('registrar_venda')

        try:
            with transaction.atomic():
                # ── VALIDAÇÃO DE STOCK ANTES DE CRIAR A VENDA ──
                erros_stock = []
                produtos_venda = []

                for i in itens:
                    # select_for_update bloqueia o registo durante a transação
                    # evita que duas vendas simultâneas esgotem o mesmo stock
                    prod = Produto.objects.select_for_update().get(id=i['id'])
                    qtd_pedida = int(i['quantidade'])

                    if prod.stock_actual < qtd_pedida:
                        erros_stock.append(
                            f"{prod.nome} — stock disponível: {prod.stock_actual} un. (pedido: {qtd_pedida} un.)"
                        )
                    else:
                        produtos_venda.append((prod, qtd_pedida))

                # Se houver qualquer erro de stock, cancela tudo e avisa
                if erros_stock:
                    for erro in erros_stock:
                        messages.error(request, f"Stock insuficiente: {erro}")
                    return redirect('registrar_venda')

                # Tudo ok — criar a venda
                venda = Venda.objects.create(
                    utilizador=request.user,
                    metodo_pagamento=request.POST.get('metodo_pagamento')
                )
                for prod, qtd in produtos_venda:
                    ItemVenda.objects.create(
                        venda=venda,
                        produto=prod,
                        quantidade=qtd,
                        preco_unitario=prod.preco_venda
                    )

            messages.success(request, f"Venda #{venda.id} registada com sucesso!")
            return redirect('ver_fatura', venda_id=venda.id)

        except Exception as e:
            messages.error(request, f"Erro ao registar a venda. Tenta novamente.")
            return redirect('registrar_venda')

    return render(request, 'venda.html', {'produtos': Produto.objects.filter(stock_actual__gt=0)})


@login_required
def extrato_caixa(request):
    # Filtros
    tipo_filtro = request.GET.get('tipo', '')
    data_inicio = request.GET.get('data_inicio', '')
    data_fim = request.GET.get('data_fim', '')

    movs = []

    if not tipo_filtro or tipo_filtro == 'entrada':
        for v in Venda.objects.all():
            movs.append({'id': f"V-{v.id}", 'data': v.data.date(), 'raw_id': v.id, 'desc': "Venda", 'tipo': 'Entrada', 'valor': float(v.valor_total), 'cor': 'text-success'})
        for e in ReceitaExtra.objects.all():
            movs.append({'id': f"R-{e.id}", 'data': e.data, 'raw_id': e.id, 'desc': e.descricao, 'tipo': 'Entrada', 'valor': float(e.valor), 'cor': 'text-success'})

    if not tipo_filtro or tipo_filtro == 'saida':
        for d in Despesa.objects.all():
            movs.append({'id': f"D-{d.id}", 'data': d.data, 'raw_id': d.id, 'desc': d.descricao, 'tipo': 'Saída', 'valor': float(d.valor), 'cor': 'text-danger'})
        for c in Compra.objects.all():
            movs.append({'id': f"C-{c.id}", 'data': c.data.date(), 'raw_id': c.id, 'desc': "Compra", 'tipo': 'Saída', 'valor': float(c.valor_total), 'cor': 'text-danger'})

    # Filtro por data
    from datetime import date
    if data_inicio:
        movs = [m for m in movs if m['data'] >= date.fromisoformat(data_inicio)]
    if data_fim:
        movs = [m for m in movs if m['data'] <= date.fromisoformat(data_fim)]

    movimentacoes = sorted(movs, key=lambda x: (x['data'], x['raw_id']), reverse=True)

    t_e = sum(m['valor'] for m in movimentacoes if m['tipo'] == 'Entrada')
    t_s = sum(m['valor'] for m in movimentacoes if m['tipo'] == 'Saída')
    paginator = Paginator(movimentacoes, 25)
    page = paginator.get_page(request.GET.get('page'))

    return render(request, 'extrato.html', {
    'movimentacoes': page,
    'saldo_final': t_e - t_s,
    'total_entradas': t_e,
    'total_saidas': t_s,
    })

@login_required
def lista_produtos(request):
    query = request.GET.get('q')
    produtos = Produto.objects.filter(Q(nome__icontains=query) | Q(marca__icontains=query)) if query else Produto.objects.all()
    return render(request, 'produtos.html', {'produtos': produtos, 'query': query})


@login_required
def planeamento_compras(request):
    produtos = Produto.objects.all().order_by('nome')
    sugestoes = [p for p in produtos if p.stock_actual <= p.stock_minimo or p.classe_abc() == 'A']
    erro = request.GET.get('erro')
    return render(request, 'planeamento.html', {'produtos': produtos, 'sugestoes': sugestoes, 'erro': erro})


@login_required
def relatorios(request):
    from django.utils import timezone
    from django.db.models import Count

    # Filtro por ano
    ano_selecionado = request.GET.get('ano')
    anos_disponiveis = sorted(set(
        list(Venda.objects.dates('data', 'year').values_list('data__year', flat=True)) +
        list(Compra.objects.dates('data', 'year').values_list('data__year', flat=True))
    ), reverse=True)

    # Recolher todos os meses com atividade
    filtro_ano = {'data__year': ano_selecionado} if ano_selecionado else {}
    filtro_ano_date = {'data__year': ano_selecionado} if ano_selecionado else {}

    meses_vendas = set(Venda.objects.filter(**filtro_ano).dates('data', 'month'))
    meses_despesas = set(Despesa.objects.filter(**filtro_ano_date).dates('data', 'month'))
    meses_compras = set(Compra.objects.filter(**filtro_ano).dates('data', 'month'))
    meses_extras = set(ReceitaExtra.objects.filter(**filtro_ano_date).dates('data', 'month'))
    todos_meses = sorted(meses_vendas | meses_despesas | meses_compras | meses_extras, reverse=True)

    relatorio_final = []
    for d_mes in todos_meses:
        v = float(Venda.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
        e = float(ReceitaExtra.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor'))['valor__sum'] or 0)
        d = float(Despesa.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor'))['valor__sum'] or 0)
        c = float(Compra.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)

        receitas = v + e
        custos = d + c
        lucro = receitas - custos
        margem = (lucro / receitas * 100) if receitas > 0 else 0

        relatorio_final.append({
            'mes': d_mes, 'vendas': receitas,
            'saidas': custos, 'lucro': lucro, 'margem': margem,
        })

    total_receitas = sum(r['vendas'] for r in relatorio_final)
    total_custos = sum(r['saidas'] for r in relatorio_final)
    total_lucro = total_receitas - total_custos
    margem_media = (total_lucro / total_receitas * 100) if total_receitas > 0 else 0

    # TOP 5 PRODUTOS MAIS VENDIDOS
    from django.db.models import Sum as DSum
    top_produtos = (
        ItemVenda.objects
        .filter(**({'venda__data__year': ano_selecionado} if ano_selecionado else {}))
        .values('produto__nome', 'produto__marca')
        .annotate(
            total_vendido=DSum('quantidade'),
            receita=DSum(F('quantidade') * F('preco_unitario'))
        )
        .order_by('-receita')[:5]
    )

    # TOP 3 FORNECEDORES
    top_fornecedores = (
        Compra.objects
        .filter(**filtro_ano)
        .values('fornecedor__nome')
        .annotate(total_gasto=DSum('valor_total'))
        .order_by('-total_gasto')[:3]
    )

    # PRODUTOS ESTAGNADOS (sem vendas)
    ids_vendidos = ItemVenda.objects.values_list('produto_id', flat=True).distinct()
    produtos_estagnados = Produto.objects.exclude(id__in=ids_vendidos).filter(stock_actual__gt=0)[:5]

    return render(request, 'relatorios.html', {
        'relatorio_final': relatorio_final,
        'total_receitas': total_receitas,
        'total_custos': total_custos,
        'total_lucro': total_lucro,
        'margem_media': margem_media,
        'top_produtos': top_produtos,
        'top_fornecedores': top_fornecedores,
        'produtos_estagnados': produtos_estagnados,
        'anos_disponiveis': anos_disponiveis,
        'ano_selecionado': ano_selecionado,
    })

@login_required
def lista_vendas(request): return render(request, 'lista_vendas.html', {'vendas': Venda.objects.all().order_by('-data')})

@login_required
def ver_fatura(request, venda_id): return render(request, 'fatura.html', {'venda': get_object_or_404(Venda, id=venda_id)})

@login_required
def lista_fornecedores(request): return render(request, 'fornecedores.html', {'fornecedores': Fornecedor.objects.all()})

@login_required
def add_fornecedor(request):
    if request.method == "POST":
        Fornecedor.objects.create(nome=request.POST.get('nome'), contacto=request.POST.get('contacto'))
        return redirect('lista_fornecedores')
    return render(request, 'add_generic.html', {'titulo': 'Novo Fornecedor'})

@login_required
def add_despesa(request):
    if request.method == "POST":
        Despesa.objects.create(
            descricao=request.POST.get('descricao'),
            valor=request.POST.get('valor'),
            data=request.POST.get('data') or timezone.now().date()
        )
        messages.success(request, "Despesa registada com sucesso!")
        return redirect('add_despesa')
    
    despesas = Despesa.objects.all().order_by('-data')
    mes = request.GET.get('mes')
    ano = request.GET.get('ano')
    if mes:
        despesas = despesas.filter(data__month=mes)
    if ano:
        despesas = despesas.filter(data__year=ano)
    
    total = despesas.aggregate(Sum('valor'))['valor__sum'] or 0
    return render(request, 'despesas.html', {
        'despesas': despesas,
        'total': total,
        'hoje': timezone.now().date(),
        'mes_selecionado': mes,
        'ano_selecionado': ano,
    })

@login_required
def add_receita(request):
    if request.method == "POST":
        ReceitaExtra.objects.create(
            descricao=request.POST.get('descricao'),
            valor=request.POST.get('valor'),
            data=request.POST.get('data') or timezone.now().date()
        )
        messages.success(request, "Receita registada com sucesso!")
        return redirect('add_receita')
    
    receitas = ReceitaExtra.objects.all().order_by('-data')
    mes = request.GET.get('mes')
    ano = request.GET.get('ano')
    if mes:
        receitas = receitas.filter(data__month=mes)
    if ano:
        receitas = receitas.filter(data__year=ano)
    
    total = receitas.aggregate(Sum('valor'))['valor__sum'] or 0
    return render(request, 'receitas.html', {
        'receitas': receitas,
        'total': total,
        'hoje': timezone.now().date(),
        'mes_selecionado': mes,
        'ano_selecionado': ano,
    })

@login_required
def add_produto(request):
    if not request.user.is_superuser:
        return redirect('lista_produtos')

    categorias = Categoria.objects.all().order_by('nome')
    errors = {}
    form_data = {}

    if request.method == "POST":
        form_data = request.POST
        nome         = request.POST.get('nome', '').strip()
        marca        = request.POST.get('marca', '').strip()
        categoria_id = request.POST.get('categoria')
        preco_venda  = request.POST.get('preco_venda')
        stock_minimo = request.POST.get('stock_minimo', 5)

        if not nome:
            errors['nome'] = "O nome do produto é obrigatório."
        if not marca:
            errors['marca'] = "A marca é obrigatória."
        if not categoria_id:
            errors['categoria'] = "Selecciona uma categoria."
        if not preco_venda:
            errors['preco_venda'] = "O preço de venda é obrigatório."
        if nome and marca and Produto.objects.filter(nome__iexact=nome, marca__iexact=marca).exists():
            errors['duplicado'] = f"Já existe um produto '{nome}' da marca '{marca}'."

        if not errors:
            try:
                categoria = Categoria.objects.get(id=categoria_id)
                Produto.objects.create(
                    nome=nome,
                    marca=marca,
                    categoria=categoria,
                    preco_venda=preco_venda,
                    stock_minimo=int(stock_minimo),
                )
                messages.success(request, f"Produto '{nome}' criado com sucesso!")
                return redirect('lista_produtos')
            except Exception as e:
                errors['geral'] = f"Erro ao guardar: {str(e)}"

    return render(request, 'add_produto.html', {
        'titulo': 'Novo Produto',
        'categorias': categorias,
        'errors': errors,
        'form_data': form_data,
    })


@login_required
def editar_produto(request, produto_id):
    if not request.user.is_superuser:
        return redirect('lista_produtos')

    produto    = get_object_or_404(Produto, id=produto_id)
    categorias = Categoria.objects.all().order_by('nome')
    errors     = {}

    if request.method == "POST":
        nome         = request.POST.get('nome', '').strip()
        marca        = request.POST.get('marca', '').strip()
        categoria_id = request.POST.get('categoria')
        preco_venda  = request.POST.get('preco_venda')
        stock_minimo = request.POST.get('stock_minimo', 5)

        if not nome:
            errors['nome'] = "O nome do produto é obrigatório."
        if not marca:
            errors['marca'] = "A marca é obrigatória."
        if not categoria_id:
            errors['categoria'] = "Selecciona uma categoria."
        if not preco_venda:
            errors['preco_venda'] = "O preço de venda é obrigatório."
        if nome and marca and Produto.objects.filter(
            nome__iexact=nome, marca__iexact=marca
        ).exclude(id=produto_id).exists():
            errors['duplicado'] = f"Já existe outro produto '{nome}' da marca '{marca}'."

        if not errors:
            try:
                produto.nome         = nome
                produto.marca        = marca
                produto.categoria    = Categoria.objects.get(id=categoria_id)
                produto.preco_venda  = preco_venda
                produto.stock_minimo = int(stock_minimo)
                produto.save()
                messages.success(request, f"Produto '{nome}' atualizado com sucesso!")
                return redirect('lista_produtos')
            except Exception as e:
                errors['geral'] = f"Erro ao guardar: {str(e)}"

    form_data = {
        'nome':        produto.nome,
        'marca':       produto.marca,
        'categoria':   str(produto.categoria_id),
        'preco_venda': produto.preco_venda,
        'stock_minimo': produto.stock_minimo,
    }

    return render(request, 'add_produto.html', {
        'titulo': f'Editar Produto — {produto.nome}',
        'categorias': categorias,
        'errors': errors,
        'form_data': form_data,
        'produto': produto,
    })

@login_required
def lista_compras(request):
    compras = Compra.objects.all().order_by('-data')
    fornecedores = Fornecedor.objects.all().order_by('nome')

    # Filtros
    fornecedor_id = request.GET.get('fornecedor')
    data_inicio = request.GET.get('data_inicio')
    data_fim = request.GET.get('data_fim')

    if fornecedor_id:
        compras = compras.filter(fornecedor__id=fornecedor_id)
    if data_inicio:
        compras = compras.filter(data__date__gte=data_inicio)
    if data_fim:
        compras = compras.filter(data__date__lte=data_fim)

    total_gasto = sum(c.valor_total for c in compras)

    return render(request, 'lista_compras.html', {
        'compras': compras,
        'fornecedores': fornecedores,
        'total_gasto': total_gasto,
    })

@login_required
def add_compra(request):
    if not request.user.is_superuser:
        return redirect('dashboard')

    produtos     = Produto.objects.all().order_by('nome')
    fornecedores = Fornecedor.objects.all().order_by('nome')
    erro = None

    if request.method == "POST":
        fornecedor_id = request.POST.get('fornecedor')
        itens_json    = request.POST.get('itens_dados', '[]')

        try:
            itens = json.loads(itens_json)
        except Exception:
            itens = []

        if not fornecedor_id:
            erro = "Selecciona um fornecedor."
        elif not itens:
            erro = "Adiciona pelo menos um produto à compra."
        else:
            try:
                with transaction.atomic():
                    fornecedor = Fornecedor.objects.get(id=fornecedor_id)
                    compra = Compra.objects.create(fornecedor=fornecedor)
                    for i in itens:
                        produto = Produto.objects.get(id=i['id'])
                        ItemCompra.objects.create(
                            compra=compra,
                            produto=produto,
                            quantidade=int(i['quantidade']),
                            preco_custo=i['preco_custo'],
                            validade=i['validade'],
                            lote=i.get('lote', '') or '',
                        )
                messages.success(request, "Compra registada com sucesso!")
                return redirect('lista_produtos')
            except Exception as e:
                erro = f"Erro ao registar compra: {str(e)}"

    return render(request, 'add_compra.html', {
        'produtos': produtos,
        'fornecedores': fornecedores,
        'erro': erro,
    })


@login_required
def ajuste_stock(request):
    if not request.user.is_superuser:
        return redirect('planeamento_compras')

    if request.method == "POST":
        produto_id = request.POST.get('produto_id')
        tipo       = request.POST.get('tipo')
        quantidade = int(request.POST.get('quantidade', 0))

        try:
            produto = Produto.objects.get(id=produto_id)
            if tipo == 'remover':
                if quantidade > produto.stock_actual:
                    messages.error(request, f"Quantidade superior ao stock actual ({produto.stock_actual} un.).")
                    return redirect('planeamento_compras')
                produto.stock_actual -= quantidade
            else:
                produto.stock_actual += quantidade
            produto.save()
            messages.success(request, f"Stock de '{produto.nome}' ajustado com sucesso.")
        except Exception:
            messages.error(request, "Erro ao ajustar stock. Tenta novamente.")

    return redirect('planeamento_compras')

@login_required
def exportar_compras_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="historico_compras.csv"'
    response.write('\ufeff')  # BOM para Excel abrir corretamente

    writer = csv.writer(response)
    writer.writerow(['ID', 'Data', 'Fornecedor', 'Produto', 'Quantidade', 'Custo Unit.', 'Validade', 'Lote', 'Total Compra'])

    compras = Compra.objects.all().order_by('-data')
    for c in compras:
        for item in c.itens.all():
            writer.writerow([
                f'#{c.id}',
                c.data.strftime('%d/%m/%Y %H:%M'),
                c.fornecedor.nome if c.fornecedor else '—',
                item.produto.nome,
                item.quantidade,
                item.preco_custo,
                item.validade.strftime('%d/%m/%Y'),
                item.lote or '—',
                c.valor_total,
            ])

    return response

@login_required
def exportar_vendas_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="historico_vendas.csv"'
    response.write('\ufeff')

    writer = csv.writer(response)
    writer.writerow(['ID', 'Data', 'Utilizador', 'Produto', 'Quantidade', 'Preço Unit.', 'Subtotal', 'Método Pagamento', 'Total Venda'])

    for v in Venda.objects.all().order_by('-data'):
        for item in v.itens.all():
            writer.writerow([
                f'#{v.id}',
                v.data.strftime('%d/%m/%Y %H:%M'),
                v.utilizador.username,
                item.produto.nome,
                item.quantidade,
                item.preco_unitario,
                item.total_item(),
                v.get_metodo_pagamento_display(),
                v.valor_total,
            ])

    return response


@login_required  
def exportar_extrato_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="extrato_caixa.csv"'
    response.write('\ufeff')

    writer = csv.writer(response)
    writer.writerow(['ID', 'Data', 'Descrição', 'Tipo', 'Valor (Kz)'])

    for v in Venda.objects.all():
        writer.writerow([f'V-{v.id}', v.data.strftime('%d/%m/%Y'), 'Venda', 'Entrada', v.valor_total])
    for e in ReceitaExtra.objects.all():
        writer.writerow([f'R-{e.id}', e.data.strftime('%d/%m/%Y'), e.descricao, 'Entrada', e.valor])
    for d in Despesa.objects.all():
        writer.writerow([f'D-{d.id}', d.data.strftime('%d/%m/%Y'), d.descricao, 'Saída', d.valor])
    for c in Compra.objects.all():
        writer.writerow([f'C-{c.id}', c.data.strftime('%d/%m/%Y'), 'Compra de stock', 'Saída', c.valor_total])

    return response

# views.py
@login_required
def exportar_relatorio_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="relatorios_mensais.csv"'
    response.write('\ufeff')

    writer = csv.writer(response)
    writer.writerow(['Mês/Ano', 'Receitas (Kz)', 'Custos (Kz)', 'Resultado (Kz)', 'Margem (%)'])

    meses_vendas = set(Venda.objects.dates('data', 'month'))
    meses_despesas = set(Despesa.objects.dates('data', 'month'))
    meses_compras = set(Compra.objects.dates('data', 'month'))
    meses_extras = set(ReceitaExtra.objects.dates('data', 'month'))
    todos_meses = sorted(meses_vendas | meses_despesas | meses_compras | meses_extras, reverse=True)

    for d_mes in todos_meses:
        v = float(Venda.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
        e = float(ReceitaExtra.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor'))['valor__sum'] or 0)
        d = float(Despesa.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor'))['valor__sum'] or 0)
        c = float(Compra.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
        receitas = v + e
        custos = d + c
        lucro = receitas - custos
        margem = (lucro / receitas * 100) if receitas > 0 else 0
        writer.writerow([d_mes.strftime('%B %Y'), receitas, custos, lucro, f"{margem:.1f}"])

    return response

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Pt, RGBColor
import io

# ─── HELPERS ──────────────────────────────────────────────────────────────

def pdf_response(filename):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def word_response(filename):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def estilo_tabela_pdf():
    return TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#4b0082')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 9),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#faf5ff')]),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e0d0f0')),
        ('FONTSIZE', (0,1), (-1,-1), 8),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ])

def cabecalho_word(doc, titulo):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading(titulo, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x4b, 0x00, 0x82)
    doc.add_paragraph(f"Universo de Beleza — Exportado em {timezone.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()

# ─── VENDAS PDF ───────────────────────────────────────────────────────────

@login_required
def exportar_vendas_pdf(request):
    response = pdf_response('historico_vendas.pdf')
    doc = SimpleDocTemplate(response, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    elementos = []

    elementos.append(Paragraph('Histórico de Vendas', styles['Title']))
    elementos.append(Paragraph(f'Universo de Beleza — {timezone.now().strftime("%d/%m/%Y %H:%M")}', styles['Normal']))
    elementos.append(Spacer(1, 0.5*cm))

    dados = [['ID', 'Data', 'Utilizador', 'Método', 'Total (Kz)']]
    for v in Venda.objects.all().order_by('-data'):
        dados.append([
            f'#{v.id}',
            v.data.strftime('%d/%m/%Y %H:%M'),
            v.utilizador.username,
            v.get_metodo_pagamento_display(),
            f'{v.valor_total:,.2f}',
        ])

    tabela = Table(dados, colWidths=[2*cm, 4*cm, 3*cm, 3*cm, 4*cm])
    tabela.setStyle(estilo_tabela_pdf())
    elementos.append(tabela)
    doc.build(elementos)
    return response

# ─── VENDAS WORD ──────────────────────────────────────────────────────────

@login_required
def exportar_vendas_word(request):
    doc = Document()
    cabecalho_word(doc, 'Histórico de Vendas')
    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = 'Table Grid'
    cabecalhos = ['ID', 'Data', 'Utilizador', 'Método', 'Total (Kz)']
    for i, h in enumerate(cabecalhos):
        cell = tabela.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for v in Venda.objects.all().order_by('-data'):
        row = tabela.add_row().cells
        row[0].text = f'#{v.id}'
        row[1].text = v.data.strftime('%d/%m/%Y %H:%M')
        row[2].text = v.utilizador.username
        row[3].text = v.get_metodo_pagamento_display()
        row[4].text = f'{v.valor_total:,.2f} Kz'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = word_response('historico_vendas.docx')
    response.write(buffer.read())
    return response

# ─── COMPRAS PDF ──────────────────────────────────────────────────────────

@login_required
def exportar_compras_pdf(request):
    response = pdf_response('historico_compras.pdf')
    doc = SimpleDocTemplate(response, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    elementos = []

    elementos.append(Paragraph('Histórico de Compras', styles['Title']))
    elementos.append(Paragraph(f'Universo de Beleza — {timezone.now().strftime("%d/%m/%Y %H:%M")}', styles['Normal']))
    elementos.append(Spacer(1, 0.5*cm))

    dados = [['ID', 'Data', 'Fornecedor', 'Produto', 'Qtd', 'Custo Unit.', 'Validade']]
    for c in Compra.objects.all().order_by('-data'):
        for item in c.itens.all():
            dados.append([
                f'#{c.id}',
                c.data.strftime('%d/%m/%Y'),
                c.fornecedor.nome if c.fornecedor else '—',
                item.produto.nome,
                str(item.quantidade),
                f'{item.preco_custo:,.2f}',
                item.validade.strftime('%d/%m/%Y'),
            ])

    tabela = Table(dados, colWidths=[1.5*cm, 2.5*cm, 3*cm, 3.5*cm, 1.5*cm, 2.5*cm, 2.5*cm])
    tabela.setStyle(estilo_tabela_pdf())
    elementos.append(tabela)
    doc.build(elementos)
    return response

# ─── COMPRAS WORD ─────────────────────────────────────────────────────────

@login_required
def exportar_compras_word(request):
    doc = Document()
    cabecalho_word(doc, 'Histórico de Compras')
    tabela = doc.add_table(rows=1, cols=7)
    tabela.style = 'Table Grid'
    cabecalhos = ['ID', 'Data', 'Fornecedor', 'Produto', 'Qtd', 'Custo Unit.', 'Validade']
    for i, h in enumerate(cabecalhos):
        cell = tabela.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for c in Compra.objects.all().order_by('-data'):
        for item in c.itens.all():
            row = tabela.add_row().cells
            row[0].text = f'#{c.id}'
            row[1].text = c.data.strftime('%d/%m/%Y')
            row[2].text = c.fornecedor.nome if c.fornecedor else '—'
            row[3].text = item.produto.nome
            row[4].text = str(item.quantidade)
            row[5].text = f'{item.preco_custo:,.2f} Kz'
            row[6].text = item.validade.strftime('%d/%m/%Y')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = word_response('historico_compras.docx')
    response.write(buffer.read())
    return response

# ─── EXTRATO PDF ──────────────────────────────────────────────────────────

@login_required
def exportar_extrato_pdf(request):
    response = pdf_response('extrato_caixa.pdf')
    doc = SimpleDocTemplate(response, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    elementos = []

    elementos.append(Paragraph('Extrato de Caixa', styles['Title']))
    elementos.append(Paragraph(f'Universo de Beleza — {timezone.now().strftime("%d/%m/%Y %H:%M")}', styles['Normal']))
    elementos.append(Spacer(1, 0.5*cm))

    movs = []
    for v in Venda.objects.all(): movs.append({'data': v.data.date(), 'desc': 'Venda', 'tipo': 'Entrada', 'valor': float(v.valor_total)})
    for e in ReceitaExtra.objects.all(): movs.append({'data': e.data, 'desc': e.descricao, 'tipo': 'Entrada', 'valor': float(e.valor)})
    for d in Despesa.objects.all(): movs.append({'data': d.data, 'desc': d.descricao, 'tipo': 'Saída', 'valor': float(d.valor)})
    for c in Compra.objects.all(): movs.append({'data': c.data.date(), 'desc': 'Compra', 'tipo': 'Saída', 'valor': float(c.valor_total)})
    movs = sorted(movs, key=lambda x: x['data'], reverse=True)

    dados = [['Data', 'Descrição', 'Tipo', 'Valor (Kz)']]
    for m in movs:
        dados.append([
            m['data'].strftime('%d/%m/%Y'),
            m['desc'],
            m['tipo'],
            f'{m["valor"]:,.2f}',
        ])

    tabela = Table(dados, colWidths=[3*cm, 7*cm, 3*cm, 4*cm])
    tabela.setStyle(estilo_tabela_pdf())
    elementos.append(tabela)
    doc.build(elementos)
    return response

# ─── EXTRATO WORD ─────────────────────────────────────────────────────────

@login_required
def exportar_extrato_word(request):
    doc = Document()
    cabecalho_word(doc, 'Extrato de Caixa')
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    for i, h in enumerate(['Data', 'Descrição', 'Tipo', 'Valor (Kz)']):
        cell = tabela.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    movs = []
    for v in Venda.objects.all(): movs.append({'data': v.data.date(), 'desc': 'Venda', 'tipo': 'Entrada', 'valor': float(v.valor_total)})
    for e in ReceitaExtra.objects.all(): movs.append({'data': e.data, 'desc': e.descricao, 'tipo': 'Entrada', 'valor': float(e.valor)})
    for d in Despesa.objects.all(): movs.append({'data': d.data, 'desc': d.descricao, 'tipo': 'Saída', 'valor': float(d.valor)})
    for c in Compra.objects.all(): movs.append({'data': c.data.date(), 'desc': 'Compra', 'tipo': 'Saída', 'valor': float(c.valor_total)})
    movs = sorted(movs, key=lambda x: x['data'], reverse=True)

    for m in movs:
        row = tabela.add_row().cells
        row[0].text = m['data'].strftime('%d/%m/%Y')
        row[1].text = m['desc']
        row[2].text = m['tipo']
        row[3].text = f'{m["valor"]:,.2f} Kz'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = word_response('extrato_caixa.docx')
    response.write(buffer.read())
    return response

# ─── RELATÓRIOS PDF ───────────────────────────────────────────────────────

@login_required
def exportar_relatorio_pdf(request):
    response = pdf_response('relatorios_mensais.pdf')
    doc = SimpleDocTemplate(response, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    elementos = []

    elementos.append(Paragraph('Relatórios Mensais', styles['Title']))
    elementos.append(Paragraph(f'Universo de Beleza — {timezone.now().strftime("%d/%m/%Y %H:%M")}', styles['Normal']))
    elementos.append(Spacer(1, 0.5*cm))

    meses_vendas = set(Venda.objects.dates('data', 'month'))
    meses_despesas = set(Despesa.objects.dates('data', 'month'))
    meses_compras = set(Compra.objects.dates('data', 'month'))
    meses_extras = set(ReceitaExtra.objects.dates('data', 'month'))
    todos_meses = sorted(meses_vendas | meses_despesas | meses_compras | meses_extras, reverse=True)

    dados = [['Mês / Ano', 'Receitas (Kz)', 'Custos (Kz)', 'Resultado (Kz)', 'Margem %']]
    for d_mes in todos_meses:
        v = float(Venda.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
        e = float(ReceitaExtra.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor'))['valor__sum'] or 0)
        d = float(Despesa.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor'))['valor__sum'] or 0)
        c = float(Compra.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
        receitas = v + e
        custos = d + c
        lucro = receitas - custos
        margem = (lucro / receitas * 100) if receitas > 0 else 0
        dados.append([
            d_mes.strftime('%B %Y').upper(),
            f'{receitas:,.2f}',
            f'{custos:,.2f}',
            f'{lucro:,.2f}',
            f'{margem:.1f}%',
        ])

    tabela = Table(dados, colWidths=[4*cm, 3.5*cm, 3.5*cm, 3.5*cm, 2.5*cm])
    tabela.setStyle(estilo_tabela_pdf())
    elementos.append(tabela)
    doc.build(elementos)
    return response

# ─── RELATÓRIOS WORD ──────────────────────────────────────────────────────

@login_required
def exportar_relatorio_word(request):
    doc = Document()
    cabecalho_word(doc, 'Relatórios Mensais')
    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = 'Table Grid'
    for i, h in enumerate(['Mês / Ano', 'Receitas (Kz)', 'Custos (Kz)', 'Resultado (Kz)', 'Margem %']):
        cell = tabela.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    meses_vendas = set(Venda.objects.dates('data', 'month'))
    meses_despesas = set(Despesa.objects.dates('data', 'month'))
    meses_compras = set(Compra.objects.dates('data', 'month'))
    meses_extras = set(ReceitaExtra.objects.dates('data', 'month'))
    todos_meses = sorted(meses_vendas | meses_despesas | meses_compras | meses_extras, reverse=True)

    for d_mes in todos_meses:
        v = float(Venda.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
        e = float(ReceitaExtra.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor'))['valor__sum'] or 0)
        d = float(Despesa.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor'))['valor__sum'] or 0)
        c = float(Compra.objects.filter(data__month=d_mes.month, data__year=d_mes.year).aggregate(Sum('valor_total'))['valor_total__sum'] or 0)
        receitas = v + e
        custos = d + c
        lucro = receitas - custos
        margem = (lucro / receitas * 100) if receitas > 0 else 0
        row = tabela.add_row().cells
        row[0].text = d_mes.strftime('%B %Y').upper()
        row[1].text = f'{receitas:,.2f} Kz'
        row[2].text = f'{custos:,.2f} Kz'
        row[3].text = f'{lucro:,.2f} Kz'
        row[4].text = f'{margem:.1f}%'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = word_response('relatorios_mensais.docx')
    response.write(buffer.read())
    return response